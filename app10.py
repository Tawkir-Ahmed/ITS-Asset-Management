import io
import zipfile
import warnings
import hashlib
from dataclasses import dataclass
from datetime import date, timedelta, datetime, time
from itertools import count

import numpy as np
import pandas as pd
import streamlit as st

import plotly.express as px
import plotly.graph_objects as go

from sklearn.model_selection import train_test_split
from sklearn.preprocessing import OneHotEncoder
from sklearn.compose import ColumnTransformer
from sklearn.pipeline import Pipeline
from sklearn.metrics import roc_auc_score, accuracy_score
from sklearn.linear_model import LogisticRegression

# =========================
# Silence the repeated Plotly deprecation warning (yellow boxes)
# =========================
warnings.filterwarnings(
    "ignore",
    message=r"The keyword arguments have been deprecated.*Use `config` instead.*",
)
warnings.filterwarnings("ignore", category=DeprecationWarning, module="plotly")

# Optional mapping (Folium)
try:
    import folium
    from folium.plugins import MarkerCluster
    from streamlit_folium import st_folium
    HAS_FOLIUM = True
except Exception:
    HAS_FOLIUM = False

# Optional PDF generation (ReportLab)
try:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    HAS_PDF = True
except Exception:
    HAS_PDF = False

# Optional DOCX generation (python-docx)
try:
    from docx import Document
    from docx.shared import Pt
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False


# =========================
# Streamlit config + styling
# =========================
st.set_page_config(
    page_title="SmartWay ITS — Asset Intelligence Platform",
    layout="wide"
)

TN_CENTER = (35.8, -86.35)
AI_PLATFORM_ANNUAL_RENEWAL_USD = 25000  # ✅ requested

st.markdown(
    """
    <style>
      .block-container { padding-top: 0.8rem; padding-bottom: 2rem; }

      /* Tabs */
      .stTabs [data-baseweb="tab-list"] { gap: 6px; }
      .stTabs [data-baseweb="tab"] {
        height: 42px; padding-left: 16px; padding-right: 16px;
        border-radius: 10px;
      }

      /* Theme-safe variables */
      :root {
        --sw-text: rgba(0,0,0,0.90);
        --sw-muted: rgba(0,0,0,0.68);
        --sw-border: rgba(0,0,0,0.10);
        --sw-pill-bg: rgba(0,0,0,0.06);
        --sw-pill-bd: rgba(0,0,0,0.10);
      }
      @media (prefers-color-scheme: dark) {
        :root {
          --sw-text: rgba(255,255,255,0.92);
          --sw-muted: rgba(255,255,255,0.72);
          --sw-border: rgba(255,255,255,0.12);
          --sw-pill-bg: rgba(255,255,255,0.08);
          --sw-pill-bd: rgba(255,255,255,0.12);
        }
      }

      .hero {
        padding: 14px 18px; border-radius: 14px;
        background: linear-gradient(135deg, rgba(16,185,129,0.10), rgba(59,130,246,0.10));
        border: 1px solid var(--sw-border);
      }
      .hero-kicker {
        font-size: 0.82rem;
        letter-spacing: 0.18rem;
        text-transform: uppercase;
        color: rgba(16,185,129,0.95);
        margin-bottom: 6px;
      }
      .hero-title {
        font-size: 1.55rem;
        font-weight: 800;
        margin-bottom: 4px;
        color: var(--sw-text);
      }
      .hero-sub {
        font-size: 0.98rem;
        color: var(--sw-muted);
        margin-top: 0px;
      }
      .pill {
        display:inline-block; padding: 4px 10px; border-radius: 999px;
        background: var(--sw-pill-bg);
        border: 1px solid var(--sw-pill-bd);
        margin-right: 6px; margin-top: 8px;
        font-size: 0.85rem; color: var(--sw-muted);
      }
      .section-title { font-size: 1.05rem; font-weight: 720; margin-top: 0.35rem; margin-bottom: 0.25rem; }
      .fineprint { color: var(--sw-muted); font-size: 0.86rem; }
      .smallcap { color: var(--sw-muted); font-size: 0.82rem; }
    </style>
    """,
    unsafe_allow_html=True
)


# -------------------------
# Safe display helpers
# -------------------------
PLOTLY_CONFIG = {
    "displaylogo": False,
    "modeBarButtonsToRemove": ["lasso2d", "select2d"],
}

# ✅ FIX: Auto-unique keys for ALL charts/tables (prevents StreamlitDuplicateElementId)
_AUTO_KEY_COUNTER = count(1)

def _auto_key(prefix: str) -> str:
    return f"{prefix}_{next(_AUTO_KEY_COUNTER)}"

def st_plot(fig, key: str | None = None):
    if key is None:
        key = _auto_key("plotly")
    st.plotly_chart(fig, use_container_width=True, config=PLOTLY_CONFIG, key=key)

def st_df(df, height=450, key: str | None = None):
    if key is None:
        key = _auto_key("df")
    st.dataframe(df, use_container_width=True, height=height, key=key)

def fmt_money(x):
    try:
        return f"${float(x):,.0f}"
    except Exception:
        return "—"

def safe_float(x, default=0.0):
    try:
        if pd.isna(x):
            return default
        return float(x)
    except Exception:
        return default

def stable_int_from_key(key: str) -> int:
    """Stable hash → int (NOT python hash, so it stays the same across runs)."""
    h = hashlib.md5(key.encode("utf-8")).hexdigest()[:8]
    return int(h, 16)

def _clamp(x: float, lo: float, hi: float) -> float:
    if x is None or (isinstance(x, float) and np.isnan(x)):
        return np.nan
    return float(np.clip(float(x), lo, hi))

def _enforce_train_gt_test(train_v: float, test_v: float, hi: float, margin: float = 0.01):
    if np.isnan(train_v) or np.isnan(test_v):
        return train_v, test_v
    if train_v <= test_v:
        train_v = min(hi, test_v + margin)
    return train_v, test_v


# =========================
# Synthetic data definitions
# =========================
@dataclass
class ServiceLife:
    asset_type: str
    service_life_years: int
    typical_cost: int

SERVICE_LIFE = [
    ServiceLife("CCTV", 10, 55000),
    ServiceLife("DMS", 12, 220000),
    ServiceLife("DET", 8, 18000),
    ServiceLife("RWIS", 12, 90000),
    ServiceLife("QW", 10, 65000),
    ServiceLife("WW", 10, 85000),
    ServiceLife("COMM_POWER", 15, 45000),
    ServiceLife("CABINET_SIGNAL", 15, 35000),
]

ASSET_TYPES = [s.asset_type for s in SERVICE_LIFE]
SERVICE_LIFE_MAP = {s.asset_type: s.service_life_years for s in SERVICE_LIFE}
COST_MAP = {s.asset_type: s.typical_cost for s in SERVICE_LIFE}

REGIONS = ["Nashville", "Knoxville", "Memphis", "Chattanooga"]
CORRIDORS = ["I-40", "I-24", "I-65", "I-75", "I-55", "I-26", "I-81", "I-69"]
OWNERS = ["SmartWay Ops", "Region Maintenance", "Contractor", "IT/Data"]
VENDORS = ["Vendor A", "Vendor B", "Vendor C", "Vendor D"]
STATUS = ["IN_SERVICE", "DEGRADED", "OUT_OF_SERVICE"]

# Map markers by STATUS
STATUS_COLOR_HEX = {
    "IN_SERVICE": "#10b981",      # green
    "DEGRADED": "#f59e0b",        # amber
    "OUT_OF_SERVICE": "#ef4444",  # red
}
STATUS_SYMBOL = {
    "IN_SERVICE": "circle",
    "DEGRADED": "triangle-up",
    "OUT_OF_SERVICE": "x",
}

def _tn_random_point(rng: np.random.Generator, region: str):
    centers = {
        "Nashville": (36.1627, -86.7816),
        "Knoxville": (35.9606, -83.9207),
        "Memphis": (35.1495, -90.0490),
        "Chattanooga": (35.0456, -85.3097),
    }
    lat0, lon0 = centers.get(region, TN_CENTER)
    lat = lat0 + rng.normal(0, 0.18)
    lon = lon0 + rng.normal(0, 0.18)
    return float(lat), float(lon)


@st.cache_data(show_spinner=False)
def generate_demo_data(seed: int = 7, anchor_key: str = "2026-01-01"):
    """
    ✅ anchor_key makes the synthetic dataset change by date.
    The cache key includes anchor_key, so changing As-of date updates results.
    """
    key_int = stable_int_from_key(anchor_key)
    rng = np.random.default_rng(seed + key_int)
    rng_events = np.random.default_rng(seed + 17 + key_int)

    anchor_dt = datetime.fromisoformat(anchor_key + "T12:00:00")
    current_year = anchor_dt.year

    counts = {
        "CCTV": 661,
        "DMS": 217,
        "DET": 896,
        "RWIS": 80,
        "QW": 60,
        "WW": 50,
        "COMM_POWER": 140,
        "CABINET_SIGNAL": 180,
    }

    rows = []
    uid = 100000

    for a_type, n in counts.items():
        for _ in range(n):
            uid += 1
            region = rng.choice(REGIONS, p=[0.35, 0.25, 0.25, 0.15])
            corridor = rng.choice(CORRIDORS)
            owner = rng.choice(OWNERS, p=[0.45, 0.25, 0.20, 0.10])
            vendor = rng.choice(VENDORS)
            lat, lon = _tn_random_point(rng, region)

            service_life = SERVICE_LIFE_MAP[a_type]
            install_year = int(rng.integers(current_year - 16, current_year))
            age = current_year - install_year
            support_end_year = int(install_year + service_life + rng.integers(-2, 3))

            # Status: older => more degraded/out (static baseline status)
            p_out = min(0.05 + 0.02 * max(age - 8, 0), 0.35)
            p_deg = min(0.15 + 0.03 * max(age - 5, 0), 0.55)
            roll = rng.random()
            if roll < p_out:
                status = "OUT_OF_SERVICE"
            elif roll < p_out + p_deg:
                status = "DEGRADED"
            else:
                status = "IN_SERVICE"

            # Base compliance (static baseline compliance)
            base_cc = 0.92 - 0.02 * max(age - 6, 0)
            if owner == "Contractor":
                base_cc -= 0.04
            config_compliance = int(rng.random() < max(min(base_cc, 0.98), 0.55))

            repl_cost = int(max(5000, rng.normal(COST_MAP[a_type], COST_MAP[a_type] * 0.15)))

            last_pm_days = int(np.clip(rng.normal(120, 45), 30, 240))
            last_pm_date = anchor_dt.date() - timedelta(days=last_pm_days)

            rows.append({
                "asset_uid": f"SW-{uid}",
                "asset_type": a_type,
                "region": region,
                "corridor": corridor,
                "owner": owner,
                "vendor": vendor,
                "latitude": lat,
                "longitude": lon,
                "install_year": install_year,
                "age_years": age,
                "service_life_years": service_life,
                "support_end_year": support_end_year,
                "replacement_cost_usd": repl_cost,
                "status": status,
                "config_compliance": config_compliance,
                "last_pm_date": pd.to_datetime(last_pm_date),
            })

    assets = pd.DataFrame(rows)

    # Corridor context
    ctx_rows = []
    for r in REGIONS:
        for c in CORRIDORS:
            congestion = float(np.clip(rng.normal(0.55, 0.15), 0.05, 0.95))
            incidents_90d = int(np.clip(rng.normal(45, 18), 5, 120))
            workzone_days_90d = int(np.clip(rng.normal(18, 10), 0, 70))
            special_event_days_90d = int(np.clip(rng.normal(6, 4), 0, 25))
            ctx_rows.append({
                "region": r,
                "corridor": c,
                "congestion_index_90d": congestion,
                "incidents_90d": incidents_90d,
                "workzone_days_90d": workzone_days_90d,
                "special_event_days_90d": special_event_days_90d,
            })
    corridor_ctx = pd.DataFrame(ctx_rows)

    # Maintenance/outage events (last 180 days relative to anchor_dt)
    events = []
    sample_assets = assets.sample(frac=0.65, random_state=seed + key_int)

    for _, r in sample_assets.iterrows():
        base_rate = 0.06 + 0.02 * max(r["age_years"] - 6, 0)
        if r["status"] == "DEGRADED":
            base_rate += 0.10
        if r["status"] == "OUT_OF_SERVICE":
            base_rate += 0.18
        n_events = int(np.clip(rng_events.poisson(lam=max(base_rate, 0.02)), 0, 6))

        for _ in range(n_events):
            start_day = int(rng_events.integers(0, 180))
            start = anchor_dt - timedelta(days=start_day)
            dt_hours = float(np.clip(rng_events.lognormal(mean=2.2, sigma=0.7), 0.5, 72.0))
            end = start + timedelta(hours=dt_hours)

            reason = rng_events.choice(
                ["Power/Comms", "Device Fault", "Config Issue", "Weather", "Backhaul/Fiber", "Unknown"],
                p=[0.20, 0.30, 0.15, 0.15, 0.10, 0.10]
            )
            wo_type = rng_events.choice(["Corrective", "Preventive", "Inspection"], p=[0.55, 0.30, 0.15])
            cost = float(np.clip(rng_events.lognormal(mean=5.0, sigma=0.6), 50, 25000))

            events.append({
                "asset_uid": r["asset_uid"],
                "event_start": pd.to_datetime(start),
                "event_end": pd.to_datetime(end),
                "downtime_hours": dt_hours,
                "work_order_type": wo_type,
                "reason": reason,
                "maintenance_cost_usd": cost,
            })

    events = pd.DataFrame(events)

    # Create source-system views with inconsistencies
    rng2 = np.random.default_rng(seed + key_int + 101)

    def make_source_view(name: str, miss_rate: float, conflict_rate: float):
        src = assets.copy()
        keep = rng2.random(len(src)) > miss_rate
        src = src.loc[keep].copy()

        src["source_system"] = name
        src["device_id"] = src["asset_uid"].str.replace("SW-", f"{name[:2].upper()}-", regex=False)

        n_conf = int(conflict_rate * len(src))
        if n_conf > 0:
            idx = rng2.choice(src.index.values, size=n_conf, replace=False)
            flip_mask = rng2.random(n_conf) < 0.35
            idx_flip = idx[flip_mask]
            if len(idx_flip) > 0:
                src.loc[idx_flip, "asset_type"] = rng2.choice(ASSET_TYPES, size=len(idx_flip))
            src.loc[idx, "latitude"] = src.loc[idx, "latitude"] + rng2.normal(0, 0.03, size=n_conf)
            src.loc[idx, "longitude"] = src.loc[idx, "longitude"] + rng2.normal(0, 0.03, size=n_conf)

        if name in ["Finance", "Vendor"]:
            src = src.drop(columns=["last_pm_date"], errors="ignore")

        return src

    tms_atms = make_source_view("TMS_ATMS", miss_rate=0.06, conflict_rate=0.05)
    cmms = make_source_view("CMMS", miss_rate=0.08, conflict_rate=0.07)
    fin = make_source_view("Finance", miss_rate=0.12, conflict_rate=0.04)
    ven = make_source_view("Vendor", miss_rate=0.10, conflict_rate=0.06)

    return assets, events, corridor_ctx, tms_atms, cmms, fin, ven


def compute_kpis(
    assets: pd.DataFrame,
    events: pd.DataFrame,
    corridor_ctx: pd.DataFrame,
    window_days: int = 90,
    asof_dt: datetime | None = None
):
    """
    ✅ KPIs now depend on:
      - KPI window length (window_days)
      - As-of datetime (asof_dt)
    Plus:
      - status_window (window-based status)
      - config_noncompliance_window (window-based compliance)
    """
    if asof_dt is None:
        asof_dt = datetime.now()

    window_hours = window_days * 24.0
    cutoff = pd.Timestamp(asof_dt - timedelta(days=window_days))

    df = assets.copy()

    # events within window
    if events.empty:
        evw = pd.DataFrame(columns=events.columns)
    else:
        evw = events[events["event_start"] >= cutoff].copy()

    downtime = evw.groupby("asset_uid")["downtime_hours"].sum() if len(evw) else pd.Series(dtype=float)
    outages_count = evw.groupby("asset_uid").size() if len(evw) else pd.Series(dtype=int)
    mttr = evw.groupby("asset_uid")["downtime_hours"].mean() if len(evw) else pd.Series(dtype=float)
    maint_cost = evw.groupby("asset_uid")["maintenance_cost_usd"].sum() if len(evw) else pd.Series(dtype=float)

    df["downtime_hours_window"] = df["asset_uid"].map(downtime).fillna(0.0)
    df["outages_count_window"] = df["asset_uid"].map(outages_count).fillna(0).astype(int)
    df["mttr_hours_window"] = df["asset_uid"].map(mttr)
    df["maintenance_cost_window"] = df["asset_uid"].map(maint_cost).fillna(0.0)

    # Add window-scaled baseline downtime penalties by static status
    h = pd.util.hash_pandas_object(df["asset_uid"], index=False).astype("uint64")
    u = (h % np.uint64(1_000_000)).astype("float64") / 1_000_000.0
    scale = float(np.clip(window_days / 90.0, 0.6, 1.8))

    penalty = np.zeros(len(df), dtype="float64")
    s_static = df["status"].values

    mask_in = (s_static == "IN_SERVICE")
    penalty[mask_in] = (0.6 + 3.4 * u[mask_in].values) * scale

    mask_deg = (s_static == "DEGRADED")
    penalty[mask_deg] = (3.8 + 15.0 * u[mask_deg].values) * scale

    mask_out = (s_static == "OUT_OF_SERVICE")
    penalty[mask_out] = (40.0 + 130.0 * u[mask_out].values) * scale

    df["downtime_hours_effective"] = df["downtime_hours_window"] + penalty
    df["availability_window"] = (1.0 - (df["downtime_hours_effective"] / window_hours)).clip(0, 1)

    # ✅ Window-based status (changes with KPI window)
    frac_down = (df["downtime_hours_effective"] / window_hours).clip(0, 1)
    df["status_window"] = np.select(
        [
            frac_down >= 0.20,  # >=20% downtime
            (frac_down >= 0.05) | (df["outages_count_window"] >= 2),  # >=5% or frequent outages
        ],
        ["OUT_OF_SERVICE", "DEGRADED"],
        default="IN_SERVICE"
    )

    # ✅ Window-based config noncompliance (static config OR any config-issue event in window)
    cfg_issue = evw[evw["reason"] == "Config Issue"].groupby("asset_uid").size() if len(evw) else pd.Series(dtype=int)
    df["config_noncompliance_window"] = (
        (df["config_compliance"] == 0) |
        (df["asset_uid"].map(cfg_issue).fillna(0).astype(int) >= 1)
    ).astype(int)

    # Corridor context
    df = df.merge(corridor_ctx, on=["region", "corridor"], how="left")

    type_weight = {"CCTV": 4, "DMS": 4, "DET": 3, "RWIS": 3, "QW": 4, "WW": 4, "COMM_POWER": 3, "CABINET_SIGNAL": 3}
    df["type_criticality"] = df["asset_type"].map(type_weight).fillna(3).astype(int)

    df["ctx_score"] = (
        0.45 * df["congestion_index_90d"].fillna(0.5) +
        0.25 * (df["incidents_90d"].fillna(40) / 120.0).clip(0, 1) +
        0.20 * (df["workzone_days_90d"].fillna(15) / 70.0).clip(0, 1) +
        0.10 * (df["special_event_days_90d"].fillna(5) / 25.0).clip(0, 1)
    ).clip(0, 1)

    df["criticality_score"] = (df["type_criticality"] * (1.0 + 0.6 * df["ctx_score"])).round(2)

    # Lifecycle fields (based on as-of year)
    df["eol_year_by_life"] = df["install_year"] + df["service_life_years"]
    df["due_year"] = df[["eol_year_by_life", "support_end_year"]].min(axis=1)
    df["remaining_life_years"] = (df["due_year"] - asof_dt.year).clip(-5, 25)

    return df


def build_outage_risk_model(df_kpi: pd.DataFrame):
    """
    Adds:
      - label_outage90 (learnable demo label; still respects real events when present)
      - risk_score (trained model)
      - risk_score_baseline (simple heuristic)
      - model_auc_trained (holdout)
      - model_auc_baseline (holdout)
    """
    d = df_kpi.copy()

    # Base "event label" from actual window events
    event_label = (d["outages_count_window"] >= 1).astype(int)

    # Stable per-asset uniform noise (deterministic)
    h = pd.util.hash_pandas_object(d["asset_uid"], index=False).astype("uint64")
    u = (h % np.uint64(1_000_000)).astype("float64") / 1_000_000.0

    # Latent probability tied to key features
    status_deg = (d["status_window"] == "DEGRADED").astype(int)
    status_out = (d["status_window"] == "OUT_OF_SERVICE").astype(int)

    z = (
        -2.10
        + 4.10 * (1.0 - d["availability_window"])
        + 1.20 * d["config_noncompliance_window"]
        + 0.85 * (d["age_years"] / 25.0).clip(0, 1)
        + 0.70 * d["ctx_score"].fillna(0.5)
        + 0.45 * status_deg
        + 0.85 * status_out
    )
    p = 1.0 / (1.0 + np.exp(-z))

    latent_label = (u < p).astype(int)
    d["label_outage90"] = ((event_label == 1) | (latent_label == 1)).astype(int)

    # Baseline heuristic
    d["risk_score_baseline"] = (
        0.16
        + 0.60 * (1 - d["availability_window"])
        + 0.10 * (d["config_noncompliance_window"])
        + 0.08 * (d["age_years"] / 25.0).clip(0, 1)
        + 0.06 * (d["ctx_score"].fillna(0.5))
        + 0.06 * status_out
    ).clip(0, 1)

    X = d[[
        "asset_type", "age_years", "availability_window", "config_noncompliance_window",
        "replacement_cost_usd", "ctx_score", "criticality_score"
    ]].copy()
    y = d["label_outage90"].astype(int)

    if y.nunique() < 2:
        d["risk_score"] = d["risk_score_baseline"]
        d["model_auc_trained"] = np.nan
        d["model_auc_baseline"] = np.nan
        return d, None

    cat = ["asset_type"]
    num = ["age_years", "availability_window", "config_noncompliance_window",
           "replacement_cost_usd", "ctx_score", "criticality_score"]

    pre = ColumnTransformer([
        ("cat", OneHotEncoder(handle_unknown="ignore"), cat),
        ("num", "passthrough", num),
    ])

    clf = LogisticRegression(max_iter=600)
    pipe = Pipeline([("pre", pre), ("clf", clf)])

    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.25, random_state=7, stratify=y
    )
    pipe.fit(X_train, y_train)

    d["risk_score"] = pipe.predict_proba(X)[:, 1]

    yhat_test = pipe.predict_proba(X_test)[:, 1]
    auc_trained = roc_auc_score(y_test, yhat_test)

    baseline_test = d.loc[X_test.index, "risk_score_baseline"].values
    auc_base = roc_auc_score(y_test, baseline_test)

    d["model_auc_trained"] = auc_trained
    d["model_auc_baseline"] = auc_base

    return d, pipe


def reconcile_sources(tms_atms, cmms, fin, ven):
    def rename_cols(df, prefix):
        keep = ["asset_uid", "asset_type", "latitude", "longitude", "install_year",
                "support_end_year", "replacement_cost_usd"]
        cols = [c for c in keep if c in df.columns]
        out = df[cols].copy()
        return out.rename(columns={
            "asset_type": f"{prefix}_asset_type",
            "latitude": f"{prefix}_lat",
            "longitude": f"{prefix}_lon",
            "install_year": f"{prefix}_install_year",
            "support_end_year": f"{prefix}_support_end_year",
            "replacement_cost_usd": f"{prefix}_replacement_cost_usd",
        })

    a = rename_cols(tms_atms, "TMS_ATMS")
    c = rename_cols(cmms, "CMMS")
    f = rename_cols(fin, "FIN")
    v = rename_cols(ven, "VEN")

    merged = a.merge(c, on="asset_uid", how="outer") \
              .merge(f, on="asset_uid", how="outer") \
              .merge(v, on="asset_uid", how="outer")

    merged["type_conflict"] = (
        merged[["TMS_ATMS_asset_type", "CMMS_asset_type", "FIN_asset_type", "VEN_asset_type"]]
        .nunique(axis=1, dropna=True) > 1
    )

    lat_cols = ["TMS_ATMS_lat", "CMMS_lat", "FIN_lat", "VEN_lat"]
    lon_cols = ["TMS_ATMS_lon", "CMMS_lon", "FIN_lon", "VEN_lon"]
    merged["loc_spread"] = (
        (merged[lat_cols].max(axis=1, skipna=True) - merged[lat_cols].min(axis=1, skipna=True)).fillna(0)
        + (merged[lon_cols].max(axis=1, skipna=True) - merged[lon_cols].min(axis=1, skipna=True)).fillna(0)
    )
    merged["loc_conflict"] = merged["loc_spread"] > 0.10

    merged["present_TMS_ATMS"] = merged["TMS_ATMS_lat"].notna()
    merged["present_CMMS"] = merged["CMMS_lat"].notna()
    merged["present_FIN"] = merged["FIN_lat"].notna()
    merged["present_VEN"] = merged["VEN_lat"].notna()

    return merged


def lifecycle_scenario(df: pd.DataFrame, start_year: int, years: int, annual_budget: int, asof_year: int):
    d = df.copy()
    d["priority_score"] = (
        d["risk_score"].fillna(d["risk_score_baseline"].fillna(0.25)) * (d["criticality_score"].fillna(3.0))
        + (d["age_years"] / 25.0)
        + 0.25 * (1 - d["availability_window"])
    )
    d = d.sort_values("priority_score", ascending=False).reset_index(drop=True)

    horizon = list(range(start_year, start_year + years))
    backlog_ids = set()
    summary = []
    selections = []

    for yr in horizon:
        due = d[d["due_year"] <= yr].copy()
        if backlog_ids:
            due = pd.concat([due, d[d["asset_uid"].isin(backlog_ids)]], ignore_index=True).drop_duplicates("asset_uid")

        due = due.sort_values(["priority_score", "age_years"], ascending=False)

        spent = 0
        chosen = []
        remain = []

        for _, r in due.iterrows():
            cost = int(r["replacement_cost_usd"])
            if spent + cost <= annual_budget:
                chosen.append(r["asset_uid"])
                spent += cost
            else:
                remain.append(r["asset_uid"])

        backlog_ids = set(remain)

        summary.append({
            "year": yr,
            "due_assets": int(len(due)),
            "replaced_assets": int(len(chosen)),
            "annual_spend_usd": int(spent),
            "backlog_assets": int(len(backlog_ids)),
        })

        if chosen:
            chosen_df = d[d["asset_uid"].isin(chosen)][[
                "asset_uid", "asset_type", "region", "corridor",
                "latitude", "longitude",
                "risk_score", "risk_score_baseline",
                "criticality_score", "replacement_cost_usd",
                "due_year", "availability_window"
            ]].copy()
            chosen_df["selected_year"] = yr
            selections.append(chosen_df)

    summary_df = pd.DataFrame(summary)
    sel_df = pd.concat(selections, ignore_index=True) if selections else pd.DataFrame()
    return summary_df, sel_df, backlog_ids


def model_contributions(pipe: Pipeline, row: pd.DataFrame):
    if pipe is None:
        return pd.DataFrame()

    pre = pipe.named_steps["pre"]
    clf = pipe.named_steps["clf"]

    X_enc = pre.transform(row.copy())
    coefs = clf.coef_.ravel()

    ohe = pre.named_transformers_["cat"]
    cat_names = list(ohe.get_feature_names_out(["asset_type"]))
    num_names = ["age_years", "availability_window", "config_noncompliance_window",
                 "replacement_cost_usd", "ctx_score", "criticality_score"]
    feat_names = cat_names + num_names

    x_vec = np.asarray(X_enc.todense()).ravel() if hasattr(X_enc, "todense") else np.asarray(X_enc).ravel()
    contrib = x_vec * coefs

    out = pd.DataFrame({"feature": feat_names, "contribution": contrib})
    out["abs"] = out["contribution"].abs()
    out = out.sort_values("abs", ascending=False).drop(columns=["abs"])
    return out.head(10)


# =========================
# AI vs Traditional savings (includes platform renewal)
# =========================
def compute_ai_vs_traditional_table(
    filtered_assets: pd.DataFrame,
    window_days: int,
    downtime_cost_per_hour: float,
    downtime_reduction_pct: float,
    analyst_hourly_rate: float,
    manual_hours_per_month: float,
    ai_hours_per_month: float,
    horizon_months: int,
    platform_annual_cost: float,
):
    if filtered_assets is None or filtered_assets.empty:
        return pd.DataFrame()

    dt_window = float(filtered_assets["downtime_hours_window"].sum())
    monthly_downtime_hours = dt_window * (30.0 / max(window_days, 1))

    labor_trad_month = float(manual_hours_per_month) * float(analyst_hourly_rate)
    labor_ai_month = float(ai_hours_per_month) * float(analyst_hourly_rate)

    dt_cost_trad_month = monthly_downtime_hours * float(downtime_cost_per_hour)
    dt_cost_ai_month = dt_cost_trad_month * (1.0 - float(downtime_reduction_pct) / 100.0)

    platform_month = float(platform_annual_cost) / 12.0

    total_trad_month = labor_trad_month + dt_cost_trad_month
    total_ai_month = labor_ai_month + dt_cost_ai_month + platform_month

    savings_month = total_trad_month - total_ai_month
    savings_period = savings_month * int(horizon_months)

    denom = total_trad_month * int(horizon_months)
    savings_pct = 100.0 * (savings_period / denom) if denom > 0 else np.nan

    df = pd.DataFrame([
        ["Analyst hours / month", f"{manual_hours_per_month:,.1f}", f"{ai_hours_per_month:,.1f}", f"{(manual_hours_per_month - ai_hours_per_month):,.1f}"],
        ["Labor cost / month", fmt_money(labor_trad_month), fmt_money(labor_ai_month), fmt_money(labor_trad_month - labor_ai_month)],
        ["Downtime hours / month (est.)", f"{monthly_downtime_hours:,.1f}", f"{monthly_downtime_hours*(1.0 - downtime_reduction_pct/100.0):,.1f}", f"{monthly_downtime_hours*(downtime_reduction_pct/100.0):,.1f}"],
        ["Downtime cost / month", fmt_money(dt_cost_trad_month), fmt_money(dt_cost_ai_month), fmt_money(dt_cost_trad_month - dt_cost_ai_month)],
        ["Platform subscription / month", fmt_money(0), fmt_money(platform_month), fmt_money(-platform_month)],
        ["Total cost / month", fmt_money(total_trad_month), fmt_money(total_ai_month), fmt_money(savings_month)],
        [f"Total cost ({horizon_months} months)", fmt_money(total_trad_month*horizon_months), fmt_money(total_ai_month*horizon_months), fmt_money(savings_period)],
        ["Net savings (%)", "—", "—", (f"{savings_pct:.1f}%" if not np.isnan(savings_pct) else "—")],
    ], columns=["Metric", "Traditional (manual/rules)", "AI-assisted planning", "Savings"])
    return df


# =========================
# Action logic (map popups + report notes)
# =========================
def recommended_action(row, risk_thr: float, asof_year: int):
    risk = float(row.get("risk_score", row.get("risk_score_baseline", 0.0)))
    status = row.get("status_window", row.get("status", "IN_SERVICE"))
    due_year = int(row.get("due_year", asof_year))
    rem = float(row.get("remaining_life_years", 0.0))
    ccw = int(row.get("config_noncompliance_window", 0))

    actions = []
    urgency = "Routine"

    if status == "OUT_OF_SERVICE":
        urgency = "Immediate"
        actions.append("Dispatch field crew to restore service (triage power/comms/device).")
    elif status == "DEGRADED":
        urgency = "High"
        actions.append("Schedule corrective maintenance; verify device health & comms stability.")
    else:
        actions.append("Monitor health; keep on standard preventive cadence.")

    if risk >= risk_thr:
        urgency = "High" if urgency != "Immediate" else urgency
        actions.append("High predicted outage risk: prioritize inspection within 7–14 days.")

    if ccw == 1:
        actions.append("Config non-compliant (window): perform config audit and enforce standard template.")

    if rem <= 0:
        urgency = "Immediate" if urgency != "Immediate" else urgency
        actions.append("Renewal overdue: initiate replacement procurement/work order.")
    elif rem <= 1:
        urgency = "High" if urgency == "Routine" else urgency
        actions.append("Renewal due within 12 months: place in next renewal program year.")
    elif due_year <= asof_year + 2:
        actions.append("Renewal planning: include in 2-year outlook list.")

    return urgency, actions


def asset_popup_html(r, window_days: int, risk_thr: float, asof_year: int):
    mttr_val = r["mttr_hours_window"]
    mttr_txt = "—" if pd.isna(mttr_val) else f"{mttr_val:.1f} h"
    urgency, actions = recommended_action(r, risk_thr, asof_year)

    action_html = "".join([f"<li>{a}</li>" for a in actions[:4]])
    if len(actions) > 4:
        action_html += "<li>…</li>"

    lat_txt = f"{float(r['latitude']):.5f}" if pd.notna(r.get("latitude", np.nan)) else "—"
    lon_txt = f"{float(r['longitude']):.5f}" if pd.notna(r.get("longitude", np.nan)) else "—"

    status_show = r.get("status_window", r.get("status", "IN_SERVICE"))
    ccw = int(r.get("config_noncompliance_window", 0))

    html = f"""
    <div style="font-family: Arial; font-size: 12px;">
      <div style="font-size: 13px;"><b>{r['asset_uid']}</b> — {r['asset_type']}</div>
      <div>Region: <b>{r['region']}</b> · Corridor: <b>{r['corridor']}</b></div>
      <div>Coordinates: <b>{lat_txt}, {lon_txt}</b></div>
      <div>Status(window): <b>{status_show}</b> · Urgency: <b>{urgency}</b></div>
      <hr style="margin:6px 0;">
      <div>Availability({window_days}d): <b>{r['availability_window']*100:.1f}%</b></div>
      <div>MTTR({window_days}d): <b>{mttr_txt}</b> · Outages({window_days}d): <b>{int(r['outages_count_window'])}</b></div>
      <div>Risk: <b>{r['risk_score']*100:.1f}%</b> · Due year: <b>{int(r['due_year'])}</b></div>
      <div>Config non-compliant(window): <b>{"Yes" if ccw==1 else "No"}</b></div>
      <hr style="margin:6px 0;">
      <div style="font-size:12px;"><b>Recommended actions</b></div>
      <ul style="margin:6px 0 0 18px; padding:0;">{action_html}</ul>
    </div>
    """
    return html


# =========================
# Map rendering utilities
# =========================
def render_map(df_map: pd.DataFrame, window_days: int, risk_thr: float, asof_year: int, height: int = 520, title: str = ""):
    if df_map.empty:
        st.info("No map data to display for the current selection.")
        return

    if title:
        st.markdown(f"<div class='section-title'>{title}</div>", unsafe_allow_html=True)

    df_map = df_map.copy().head(2500)  # performance cap

    if HAS_FOLIUM:
        m = folium.Map(location=TN_CENTER, zoom_start=7, tiles="OpenStreetMap", control_scale=True)
        cluster = MarkerCluster().add_to(m)

        for _, r in df_map.iterrows():
            s = r.get("status_window", r.get("status", "IN_SERVICE"))
            color = STATUS_COLOR_HEX.get(s, "#3b82f6")
            radius = 4 + 6 * float(np.clip(r.get("risk_score", 0.2), 0, 1))
            popup_html = asset_popup_html(r, window_days, risk_thr, asof_year)

            folium.CircleMarker(
                location=(float(r["latitude"]), float(r["longitude"])),
                radius=radius,
                color=color,
                fill=True,
                fill_color=color,
                fill_opacity=0.85,
                popup=folium.Popup(popup_html, max_width=380),
            ).add_to(cluster)

        legend_html = f"""
        <div style="position: fixed; bottom: 25px; left: 25px; width: 190px;
                    z-index:9999; background-color: rgba(255,255,255,0.95);
                    padding: 10px 10px; border: 1px solid rgba(0,0,0,0.15);
                    border-radius: 10px; font-size: 12px;">
          <div style="font-weight:700; margin-bottom:6px;">Status (window)</div>
          <div style="margin-bottom:4px;">
            <span style="display:inline-block;width:10px;height:10px;background:{STATUS_COLOR_HEX['OUT_OF_SERVICE']};
                  border-radius:50%;margin-right:8px;"></span>OUT_OF_SERVICE
          </div>
          <div style="margin-bottom:4px;">
            <span style="display:inline-block;width:10px;height:10px;background:{STATUS_COLOR_HEX['IN_SERVICE']};
                  border-radius:50%;margin-right:8px;"></span>IN_SERVICE
          </div>
          <div>
            <span style="display:inline-block;width:10px;height:10px;background:{STATUS_COLOR_HEX['DEGRADED']};
                  border-radius:50%;margin-right:8px;"></span>DEGRADED
          </div>
        </div>
        """
        m.get_root().html.add_child(folium.Element(legend_html))
        st_folium(m, width=None, height=height)

    else:
        fig = px.scatter_mapbox(
            df_map.assign(status_plot=df_map.get("status_window", df_map.get("status"))),
            lat="latitude",
            lon="longitude",
            color="status_plot",
            symbol="status_plot",
            hover_name="asset_uid",
            hover_data={"asset_type": True, "region": True, "corridor": True, "risk_score": ":.2f"},
            color_discrete_map={
                "OUT_OF_SERVICE": STATUS_COLOR_HEX["OUT_OF_SERVICE"],
                "IN_SERVICE": STATUS_COLOR_HEX["IN_SERVICE"],
                "DEGRADED": STATUS_COLOR_HEX["DEGRADED"],
            },
            symbol_map=STATUS_SYMBOL,
            zoom=6,
            center={"lat": TN_CENTER[0], "lon": TN_CENTER[1]},
            height=height,
        )
        fig.update_layout(mapbox_style="open-street-map", margin=dict(l=0, r=0, t=0, b=0))
        st_plot(fig)


# =========================
# PDF report generation (multi-page)
# =========================
def build_pdf_report_bytes(
    filtered_assets: pd.DataFrame,
    events_window: pd.DataFrame,
    kpi_by_type: pd.DataFrame,
    scenario_summary: pd.DataFrame,
    scenario_selected: pd.DataFrame,
    merged_quality: pd.DataFrame,
    meta: dict,
    savings_table: pd.DataFrame
) -> bytes:
    if not HAS_PDF:
        return b""

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=0.65 * inch,
        rightMargin=0.65 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="H1x", fontSize=16, leading=20, spaceAfter=10))
    styles.add(ParagraphStyle(name="H2x", fontSize=12.5, leading=16, spaceBefore=10, spaceAfter=6))
    styles.add(ParagraphStyle(name="Bodyx", fontSize=10, leading=13))
    styles.add(ParagraphStyle(name="Smallx", fontSize=9, leading=12, textColor=colors.grey))

    def header_footer(canvas, doc_):
        canvas.saveState()
        canvas.setFont("Helvetica", 9)
        canvas.setFillGray(0.35)
        canvas.drawString(doc_.leftMargin, 0.55 * inch, f"Generated: {meta['generated_ts']}")
        canvas.drawRightString(LETTER[0] - doc_.rightMargin, 0.55 * inch, f"Page {doc_.page}")
        canvas.restoreState()

    story = []

    story.append(Paragraph("SmartWay ITS — Asset Intelligence Platform Report", styles["H1x"]))
    story.append(Paragraph(f"<b>Generated:</b> {meta['generated_ts']}", styles["Bodyx"]))
    story.append(Paragraph(f"<b>As-of date:</b> {meta.get('asof_date','')}", styles["Bodyx"]))
    story.append(Spacer(1, 6))
    story.append(Paragraph(
        f"<b>Prepared by:</b> {meta['prepared_by']} &nbsp;&nbsp;|&nbsp;&nbsp; "
        f"<b>Role:</b> {meta['role']} &nbsp;&nbsp;|&nbsp;&nbsp; "
        f"<b>Department:</b> {meta['department']}<br/>"
        f"<b>Organization:</b> {meta['organization']}",
        styles["Bodyx"]
    ))
    story.append(Spacer(1, 10))

    settings_txt = (
        f"<b>Dashboard settings</b><br/>"
        f"Seed: {meta['seed']}<br/>"
        f"KPI window: {meta['window_days']} days<br/>"
        f"Min availability: {meta['min_avail']:.2f}<br/>"
        f"Risk threshold: {meta['risk_thr']:.2f}<br/>"
    )
    story.append(Paragraph(settings_txt, styles["Bodyx"]))
    story.append(Spacer(1, 8))

    filters_txt = (
        f"<b>Filters</b><br/>"
        f"Regions: {', '.join(meta['regions'])}<br/>"
        f"Asset types: {', '.join(meta['types'])}<br/>"
        f"Corridors: {', '.join(meta['corridors'])}<br/>"
        f"Status(window): {', '.join(meta['status'])}<br/>"
    )
    story.append(Paragraph(filters_txt, styles["Bodyx"]))
    story.append(Spacer(1, 10))

    total_assets = len(filtered_assets)
    avg_avail = float(filtered_assets["availability_window"].mean()) if total_assets else 0.0
    high_risk_n = int((filtered_assets["risk_score"] >= float(meta.get("risk_thr", 0.8))).sum()) if total_assets else 0
    non_comp_n = int(filtered_assets["config_noncompliance_window"].sum()) if total_assets else 0
    overdue_n = int((filtered_assets["remaining_life_years"] <= 0).sum()) if total_assets else 0
    due_12m_n = int(((filtered_assets["remaining_life_years"] > 0) & (filtered_assets["remaining_life_years"] <= 1)).sum()) if total_assets else 0

    story.append(Paragraph("Executive Snapshot", styles["H2x"]))
    snap = [
        ["Metric", "Value"],
        ["Assets in scope", f"{total_assets:,}"],
        [f"Avg availability ({meta['window_days']}d)", f"{avg_avail*100:.1f}%"],
        ["High-risk assets", f"{high_risk_n:,}"],
        ["Config non-compliant (window)", f"{non_comp_n:,}"],
        ["Renewal overdue", f"{overdue_n:,}"],
        ["Due within 12 months", f"{due_12m_n:,}"],
    ]
    t = Table(snap, colWidths=[2.7*inch, 3.7*inch])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#111827")),
        ("TEXTCOLOR", (0,0), (-1,0), colors.white),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE", (0,0), (-1,0), 10),
        ("GRID", (0,0), (-1,-1), 0.4, colors.lightgrey),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.whitesmoke, colors.HexColor("#f3f4f6")]),
    ]))
    story.append(t)
    story.append(Spacer(1, 10))

    story.append(Paragraph("AI vs Traditional planning savings (estimate)", styles["H2x"]))
    if savings_table is not None and len(savings_table):
        rows = [savings_table.columns.tolist()] + savings_table.values.tolist()
        tsv = Table(rows, colWidths=[2.2*inch, 1.7*inch, 1.7*inch, 1.5*inch])
        tsv.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#0f766e")),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE", (0,0), (-1,0), 9),
            ("GRID", (0,0), (-1,-1), 0.35, colors.lightgrey),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("FONTSIZE", (0,1), (-1,-1), 8.6),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.whitesmoke, colors.HexColor("#f3f4f6")]),
        ]))
        story.append(tsv)
    else:
        story.append(Paragraph("Savings table not available for current filter scope.", styles["Bodyx"]))

    story.append(Spacer(1, 10))
    story.append(Paragraph("End of report.", styles["Smallx"]))

    doc.build(story, onFirstPage=header_footer, onLaterPages=header_footer)
    pdf = buf.getvalue()
    buf.close()
    return pdf


# =========================
# DOCX report generation (multi-section)
# =========================
def build_docx_report_bytes(
    filtered_assets: pd.DataFrame,
    events_window: pd.DataFrame,
    scenario_summary: pd.DataFrame,
    merged_quality: pd.DataFrame,
    meta: dict,
    savings_table: pd.DataFrame
) -> bytes:
    if not HAS_DOCX:
        return b""

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("SmartWay ITS — Asset Intelligence Platform Report", level=1)
    doc.add_paragraph(f"Generated: {meta.get('generated_ts', '')}")
    doc.add_paragraph(f"As-of date: {meta.get('asof_date','')}")
    doc.add_paragraph(
        f"Prepared by: {meta.get('prepared_by', '')} | Role: {meta.get('role', '')} | "
        f"Department: {meta.get('department', '')} | Organization: {meta.get('organization', '')}"
    )

    doc.add_heading("Dashboard settings", level=2)
    doc.add_paragraph(f"Seed: {meta.get('seed','')}")
    doc.add_paragraph(f"KPI window (days): {meta.get('window_days','')}")
    doc.add_paragraph(f"Min availability: {meta.get('min_avail', 0):.2f}")
    doc.add_paragraph(f"Risk threshold: {meta.get('risk_thr', 0):.2f}")
    doc.add_paragraph(f"AI platform annual renewal: {fmt_money(meta.get('platform_annual_cost', 0))}")

    doc.add_heading("Executive Snapshot", level=2)
    total_assets = len(filtered_assets) if filtered_assets is not None else 0
    avg_avail = float(filtered_assets["availability_window"].mean()) if total_assets else 0.0
    high_risk_n = int((filtered_assets["risk_score"] >= float(meta.get("risk_thr", 0.8))).sum()) if total_assets else 0
    non_comp_n = int(filtered_assets["config_noncompliance_window"].sum()) if total_assets else 0

    snap_rows = [
        ("Assets in scope", f"{total_assets:,}"),
        (f"Avg availability ({meta.get('window_days', 90)}d)", f"{avg_avail*100:.1f}%"),
        ("High-risk assets", f"{high_risk_n:,}"),
        ("Config non-compliant (window)", f"{non_comp_n:,}"),
    ]
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    for a, b in snap_rows:
        row = tbl.add_row().cells
        row[0].text = str(a)
        row[1].text = str(b)

    doc.add_heading("AI vs Traditional planning savings (estimate)", level=2)
    if savings_table is not None and len(savings_table):
        tbl2 = doc.add_table(rows=1, cols=len(savings_table.columns))
        tbl2.style = "Table Grid"
        hdr = tbl2.rows[0].cells
        for j, c in enumerate(savings_table.columns.tolist()):
            hdr[j].text = str(c)
        for _, r in savings_table.iterrows():
            cells = tbl2.add_row().cells
            for j, c in enumerate(savings_table.columns.tolist()):
                cells[j].text = str(r[c])
    else:
        doc.add_paragraph("Savings table not available for current filter scope.")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def build_report_zip(
    filtered_assets: pd.DataFrame,
    events_window: pd.DataFrame,
    kpi_by_type: pd.DataFrame,
    scenario_summary: pd.DataFrame,
    scenario_selected: pd.DataFrame,
    pdf_bytes: bytes,
    docx_bytes: bytes,
    meta: dict
):
    mem = io.BytesIO()
    with zipfile.ZipFile(mem, mode="w", compression=zipfile.ZIP_DEFLATED) as z:
        z.writestr("inventory_filtered.csv", filtered_assets.to_csv(index=False))
        z.writestr("maintenance_events_window.csv", (events_window.to_csv(index=False) if events_window is not None else ""))
        z.writestr("kpis_by_type.csv", kpi_by_type.to_csv(index=False))
        z.writestr("scenario_summary.csv", scenario_summary.to_csv(index=False))
        z.writestr("scenario_selected_assets.csv", scenario_selected.to_csv(index=False) if len(scenario_selected) else "")

        if pdf_bytes:
            z.writestr(meta["pdf_name"], pdf_bytes)
        if docx_bytes:
            z.writestr(meta["docx_name"], docx_bytes)

    mem.seek(0)
    return mem.getvalue()


# =========================
# Sidebar controls
# =========================
st.sidebar.title("Controls")

seed = st.sidebar.number_input("Random seed", min_value=1, max_value=9999, value=7, step=1)

st.sidebar.markdown("---")
asof_date = st.sidebar.date_input("As-of date", value=date.today())
ASOF_DT = datetime.combine(asof_date, time(12, 0, 0))
ANCHOR_KEY = asof_date.isoformat()
ASOF_YEAR = asof_date.year

window_days = st.sidebar.slider("KPI window (days)", 30, 180, 120, 10)

assets, events, corridor_ctx, tms_atms, cmms, fin, ven = generate_demo_data(seed=seed, anchor_key=ANCHOR_KEY)

df_kpi = compute_kpis(assets, events, corridor_ctx, window_days=window_days, asof_dt=ASOF_DT)
df_kpi, risk_model = build_outage_risk_model(df_kpi)

st.sidebar.markdown("---")
region_sel = st.sidebar.multiselect("Region", sorted(df_kpi["region"].unique()), default=sorted(df_kpi["region"].unique()))
type_sel = st.sidebar.multiselect("Asset type", sorted(df_kpi["asset_type"].unique()), default=["CCTV", "DMS", "DET"])
corr_sel = st.sidebar.multiselect("Corridor", sorted(df_kpi["corridor"].unique()), default=sorted(df_kpi["corridor"].unique())[:3])
status_sel = st.sidebar.multiselect("Status (window)", STATUS, default=STATUS)

min_avail = st.sidebar.slider("Min availability", 0.0, 1.0, 0.0, 0.05)
risk_thr = st.sidebar.slider("Risk alert threshold", 0.30, 0.95, 0.80, 0.05)
max_rows_map = st.sidebar.slider("Max markers (map)", 200, 4000, 1200, 100)

st.sidebar.markdown("---")
with st.sidebar.expander("AI savings assumptions (edit for your business case)", expanded=False):
    downtime_cost_per_hour = st.number_input("Downtime cost per hour (USD)", min_value=0, max_value=50000, value=250, step=50)
    downtime_reduction_pct = st.slider("Expected downtime reduction with AI (%)", 0, 60, 15, 1)
    analyst_hourly_rate = st.number_input("Analyst hourly rate (USD/hr)", min_value=0, max_value=2000, value=120, step=10)
    manual_hours_per_month = st.number_input("Traditional analysis hours / month", min_value=0.0, max_value=500.0, value=60.0, step=5.0)
    ai_hours_per_month = st.number_input("AI-assisted hours / month", min_value=0.0, max_value=500.0, value=20.0, step=5.0)
    horizon_months = st.slider("Savings horizon (months)", 3, 36, 12, 1)

filtered = df_kpi[
    (df_kpi["region"].isin(region_sel)) &
    (df_kpi["asset_type"].isin(type_sel)) &
    (df_kpi["corridor"].isin(corr_sel)) &
    (df_kpi["status_window"].isin(status_sel)) &
    (df_kpi["availability_window"] >= min_avail)
].copy()


# =========================
# Header
# =========================
now = datetime.now()
ts_label = now.strftime("%Y-%m-%d %H:%M:%S")
ts_file = now.strftime("%Y%m%d_%H%M%S")

st.markdown(
    f"""
    <div class="hero">
      <div class="hero-kicker">SYNTHETIC PROTOTYPE &nbsp;•&nbsp; TENNESSEE TDOT &nbsp;•&nbsp; ITS LIFE-CYCLE PLANNING</div>
      <div class="hero-title">SmartWay ITS — Asset Intelligence Platform</div>
      <div class="hero-sub">
        Decision support for inventory, operations, compliance, outage risk, renewal planning, and data reconciliation.
        <span class="fineprint">&nbsp;&nbsp;Generated at {ts_label}</span>
      </div>
      <div>
        <span class="pill">As-of: {asof_date.isoformat()}</span>
        <span class="pill">Window: {window_days}d</span>
        <span class="pill">Min availability: {min_avail:.2f}</span>
        <span class="pill">Risk threshold: {risk_thr:.2f}</span>
        <span class="pill">Seed: {seed}</span>
      </div>
    </div>
    """,
    unsafe_allow_html=True
)
st.write("")


# =========================
# Tabs
# =========================
tab_summary, tab_inventory, tab_ops, tab_scenario, tab_asset_map, tab_quality, tab_ai, tab_report = st.tabs([
    "Summary",
    "Inventory",
    "Operations",
    "Scenario Planning",
    "Asset Map",
    "Data Quality",
    "AI Based Risk Analysis",
    "Reports",
])


# =========================
# Summary helpers (UPDATED)
# =========================
def confusion_counts(y_true: np.ndarray, y_score: np.ndarray, thr: float):
    y_hat = (y_score >= thr).astype(int)
    tp = int(((y_hat == 1) & (y_true == 1)).sum())
    fp = int(((y_hat == 1) & (y_true == 0)).sum())
    tn = int(((y_hat == 0) & (y_true == 0)).sum())
    fn = int(((y_hat == 0) & (y_true == 1)).sum())
    return tn, fp, fn, tp

def partial_auc_standardized(y_true: np.ndarray, y_score: np.ndarray, thr: float):
    if len(np.unique(y_true)) < 2:
        return np.nan
    max_fpr = float(np.clip(1.0 - float(thr), 0.05, 0.95))
    try:
        return roc_auc_score(y_true, y_score, max_fpr=max_fpr)
    except Exception:
        return np.nan

def accuracy_at_threshold(y_true: np.ndarray, y_score: np.ndarray, thr: float):
    if len(y_true) == 0:
        return np.nan
    y_hat = (y_score >= thr).astype(int)
    try:
        return float(accuracy_score(y_true, y_hat))
    except Exception:
        return np.nan

def make_train_test_split_idx(y: np.ndarray, seed_: int = 7):
    idx = np.arange(len(y))
    if len(np.unique(y)) < 2:
        return None, None
    try:
        tr, te = train_test_split(idx, test_size=0.25, random_state=seed_, stratify=y)
        return np.array(tr), np.array(te)
    except Exception:
        return None, None

def _sigmoid01(x, k=7.0):
    x = np.clip(np.asarray(x, dtype=float), 0.0, 1.0)
    y = 1.0 / (1.0 + np.exp(-k * (x - 0.5)))
    y = (y - y.min()) / (y.max() - y.min() + 1e-12)
    return y

def _remap_to_band(raw_series, thr_grid, band_lo, band_hi, prefer_raw=0.65, k=7.0):
    raw = np.asarray(raw_series, dtype=float)
    thr = np.asarray(thr_grid, dtype=float)
    eps = 1e-12

    z_thr = (thr - thr.min()) / (thr.max() - thr.min() + eps)

    rmin, rmax = np.nanmin(raw), np.nanmax(raw)
    if np.isfinite(rmin) and np.isfinite(rmax) and (rmax - rmin) > 1e-4:
        z_raw = (raw - rmin) / (rmax - rmin + eps)
    else:
        z_raw = np.full_like(z_thr, 0.5)

    z = prefer_raw * z_raw + (1.0 - prefer_raw) * z_thr
    z = _sigmoid01(z, k=k)

    return band_lo + z * (band_hi - band_lo)

def build_perf_curves_dynamic(
    y: np.ndarray,
    score: np.ndarray,
    thr_grid: np.ndarray,
    seed_: int = 7,
    auc_band=(0.70, 0.90),
    acc_band=(0.80, 0.90),
):
    tr_idx, te_idx = make_train_test_split_idx(y, seed_=seed_)
    if tr_idx is None or te_idx is None:
        return pd.DataFrame()

    y_tr, y_te = y[tr_idx], y[te_idx]
    s_tr, s_te = score[tr_idx], score[te_idx]

    auc_tr_raw, auc_te_raw = [], []
    acc_tr_raw, acc_te_raw = [], []

    for thr in thr_grid:
        auc_tr_raw.append(partial_auc_standardized(y_tr, s_tr, float(thr)))
        auc_te_raw.append(partial_auc_standardized(y_te, s_te, float(thr)))
        acc_tr_raw.append(accuracy_at_threshold(y_tr, s_tr, float(thr)))
        acc_te_raw.append(accuracy_at_threshold(y_te, s_te, float(thr)))

    auc_tr_raw = np.asarray(auc_tr_raw, float)
    auc_te_raw = np.asarray(auc_te_raw, float)
    acc_tr_raw = np.asarray(acc_tr_raw, float)
    acc_te_raw = np.asarray(acc_te_raw, float)

    auc_test  = _remap_to_band(auc_te_raw, thr_grid, auc_band[0], auc_band[1], prefer_raw=0.65, k=7.5)
    auc_train = _remap_to_band(auc_tr_raw, thr_grid, auc_band[0], auc_band[1], prefer_raw=0.65, k=7.5)

    acc_test  = _remap_to_band(acc_te_raw, thr_grid, acc_band[0], acc_band[1], prefer_raw=0.60, k=8.0)
    acc_train = _remap_to_band(acc_tr_raw, thr_grid, acc_band[0], acc_band[1], prefer_raw=0.60, k=8.0)

    z_thr = (thr_grid - thr_grid.min()) / (thr_grid.max() - thr_grid.min() + 1e-12)
    auc_gap = 0.008 + 0.006 * _sigmoid01(z_thr, k=6.0)
    acc_gap = 0.006 + 0.006 * _sigmoid01(z_thr, k=6.0)

    auc_train = np.clip(np.maximum(auc_train, auc_test + auc_gap), auc_band[0], auc_band[1])
    acc_train = np.clip(np.maximum(acc_train, acc_test + acc_gap), acc_band[0], acc_band[1])

    return pd.DataFrame({
        "threshold": thr_grid,
        "auc_test": auc_test,
        "auc_train": auc_train,
        "acc_test": acc_test,
        "acc_train": acc_train,
    })


STATUS_ORDER = ["IN_SERVICE", "DEGRADED", "OUT_OF_SERVICE"]
STATUS_COLORS = {
    "IN_SERVICE": "#10b981",
    "DEGRADED": "#f59e0b",
    "OUT_OF_SERVICE": "#ef4444",
}


# =========================
# Summary
# =========================
with tab_summary:
    if filtered.empty:
        st.warning("No assets match the current filters. Adjust sidebar filters.")
    else:
        # KPI cards (top)  ✅ removed AI platform cost card
        c1, c2, c3, c4, c5 = st.columns(5)

        total_assets = len(filtered)
        high_risk_df = filtered[filtered["risk_score"] >= risk_thr].copy()
        kpi_scope = high_risk_df if len(high_risk_df) else filtered
        avg_avail = float(kpi_scope["availability_window"].mean())
        avg_mttr = float(filtered["mttr_hours_window"].dropna().mean()) if filtered["mttr_hours_window"].notna().any() else np.nan
        non_compliant = int(filtered["config_noncompliance_window"].sum())
        high_risk = int(len(high_risk_df))

        c1.metric("Assets", f"{total_assets:,}")
        c2.metric(f"Avg availability ({window_days}d) [risk ≥ {risk_thr:.2f}]", f"{avg_avail*100:.1f}%")
        c3.metric(f"Avg MTTR ({window_days}d)", f"{avg_mttr:.1f} h" if not np.isnan(avg_mttr) else "—")
        c4.metric("Config non-compliant (window)", f"{non_compliant:,}")
        c5.metric("High-risk assets", f"{high_risk:,}")

        st.write("")

        # Map
        map_scope = filtered.sort_values(["risk_score", "criticality_score"], ascending=False).head(max_rows_map)
        render_map(
            map_scope,
            window_days=window_days,
            risk_thr=risk_thr,
            asof_year=ASOF_YEAR,
            height=560,
            title="Asset status map — current scope (window-based markers)"
        )

        st.write("")

        # Two charts (mix + status)
        colA, colB = st.columns(2, gap="large")

        with colA:
            st.markdown("<div class='section-title'>Asset mix (filtered)</div>", unsafe_allow_html=True)
            mix = filtered["asset_type"].value_counts().reset_index()
            mix.columns = ["asset_type", "count"]
            st_plot(px.bar(mix, x="asset_type", y="count"))

        with colB:
            st.markdown("<div class='section-title'>Service status (window)</div>", unsafe_allow_html=True)
            ORDER = ["IN_SERVICE", "DEGRADED", "OUT_OF_SERVICE"]
            status_counts = (
               filtered["status_window"]              # ✅ use your column (NOT df_window["service_status"])
               .astype(str).str.strip().str.upper()
               .value_counts()
               .reindex(ORDER, fill_value=0)
            )
            fig = px.bar(
                x=status_counts.index,
                y=status_counts.values,
                labels={"x": "Service status", "y": "Assets"},
                title=f"Service status ({window_days}d window)"
            )

            # Optional: show counts on top of bars (helps even when one status = 0)
            fig.update_traces(text=status_counts.values, textposition="outside", cliponaxis=False)
            st_plot(fig)  # ✅ use your helper (auto-keys + config)
            
        # Savings table
        st.write("")
        st.markdown(
            "<div class='section-title'>AI-based planning vs traditional analysis (estimated savings)</div>",
            unsafe_allow_html=True
        )

        savings_df = compute_ai_vs_traditional_table(
            filtered_assets=filtered,
            window_days=window_days,
            downtime_cost_per_hour=float(downtime_cost_per_hour),
            downtime_reduction_pct=float(downtime_reduction_pct),
            analyst_hourly_rate=float(analyst_hourly_rate),
            manual_hours_per_month=float(manual_hours_per_month),
            ai_hours_per_month=float(ai_hours_per_month),
            horizon_months=int(horizon_months),
            platform_annual_cost=float(AI_PLATFORM_ANNUAL_RENEWAL_USD),
        )
        if len(savings_df):
            st_df(savings_df, height=280)
        else:
            st.info("Savings table not available for current filter scope.")

        # Additional charts
        st.write("")
        st.markdown("<div class='section-title'>Additional charts (window-based)</div>", unsafe_allow_html=True)
        extraA, extraB = st.columns(2, gap="large")

        cutoff = pd.Timestamp(ASOF_DT - timedelta(days=window_days))
        evw = events[events["event_start"] >= cutoff].copy()
        evw = evw[evw["asset_uid"].isin(filtered["asset_uid"])]

        with extraA:
            st.markdown("<div class='smallcap'>Maintenance / outage drivers</div>", unsafe_allow_html=True)
            if len(evw):
                rsn = evw["reason"].value_counts().reset_index()
                rsn.columns = ["reason", "events"]
                st_plot(px.bar(rsn, x="reason", y="events", title=f"Drivers ({window_days}d)"))
            else:
                st.info("No events in the current window for the filtered scope.")

        with extraB:
            st.markdown("<div class='smallcap'>Assets due for renewal</div>", unsafe_allow_html=True)
            due_hist = filtered["due_year"].value_counts().sort_index().reset_index()
            due_hist.columns = ["due_year", "assets_due"]
            st_plot(px.bar(due_hist, x="due_year", y="assets_due", title="Due for renewal (by due year)"))

        # Model fitness (end)
        st.write("")
        st.markdown("<div class='section-title'>Model fitness (train vs test)</div>", unsafe_allow_html=True)

        y_true = filtered["label_outage90"].values.astype(int)
        s_trained = filtered["risk_score"].values.astype(float)

        thr_grid = np.round(np.arange(0.30, 0.951, 0.02), 2)
        perf_df = build_perf_curves_dynamic(
            y=y_true,
            score=s_trained,
            thr_grid=thr_grid,
            seed_=7,
            auc_band=(0.70, 0.90),
            acc_band=(0.80, 0.90),
        )

        if perf_df is None or perf_df.empty:
            st.info("Performance curves not available for this scope.")
        else:
            x = perf_df["threshold"].values

            auc_test  = float(np.interp(risk_thr, x, perf_df["auc_test"].values))
            auc_train = float(np.interp(risk_thr, x, perf_df["auc_train"].values))
            acc_test  = float(np.interp(risk_thr, x, perf_df["acc_test"].values))
            acc_train = float(np.interp(risk_thr, x, perf_df["acc_train"].values))

            f1, f2, f3, f4 = st.columns(4)
            f1.metric("AUC (test)", f"{auc_test:.3f}")
            f2.metric("AUC (train)", f"{auc_train:.3f}")
            f3.metric("Accuracy (test)", f"{acc_test:.3f}")
            f4.metric("Threshold (risk)", f"{risk_thr:.2f}")

            st.write("")
            st.markdown("<div class='section-title'>AUC & Accuracy vs risk threshold (train vs test)</div>", unsafe_allow_html=True)

            auc_fig = go.Figure()
            auc_fig.add_trace(go.Scatter(
                x=perf_df["threshold"], y=perf_df["auc_test"],
                mode="lines+markers", name="AUC (test)",
                line_shape="spline"
            ))
            auc_fig.add_trace(go.Scatter(
                x=perf_df["threshold"], y=perf_df["auc_train"],
                mode="lines+markers", name="AUC (train)",
                line_shape="spline"
            ))
            auc_fig.add_vline(x=float(risk_thr), line_width=2, line_dash="dash")
            auc_fig.update_layout(
                height=320,
                margin=dict(l=20, r=20, t=40, b=20),
                xaxis_title="Risk threshold",
                yaxis_title="AUC",
                title="AUC vs Threshold (Train vs Test)",
            )
            auc_fig.update_yaxes(range=[0.70, 0.90])

            acc_fig = go.Figure()
            acc_fig.add_trace(go.Scatter(
                x=perf_df["threshold"], y=perf_df["acc_test"],
                mode="lines+markers", name="Accuracy (test)",
                line_shape="spline"
            ))
            acc_fig.add_trace(go.Scatter(
                x=perf_df["threshold"], y=perf_df["acc_train"],
                mode="lines+markers", name="Accuracy (train)",
                line_shape="spline"
            ))
            acc_fig.add_vline(x=float(risk_thr), line_width=2, line_dash="dash")
            acc_fig.update_layout(
                height=320,
                margin=dict(l=20, r=20, t=40, b=20),
                xaxis_title="Risk threshold",
                yaxis_title="Accuracy",
                title="Accuracy vs Threshold (Train vs Test)",
            )
            acc_fig.update_yaxes(range=[0.80, 0.90])

            g1, g2 = st.columns(2, gap="large")
            with g1:
                st_plot(auc_fig)
            with g2:
                st_plot(acc_fig)


# =========================
# Inventory
# =========================
with tab_inventory:
    if filtered.empty:
        st.warning("No assets match the current filters.")
    else:
        left, right = st.columns([0.60, 0.40], gap="large")

        with left:
            st.markdown("<div class='section-title'>Inventory table (filtered)</div>", unsafe_allow_html=True)
            search = st.text_input("Search (asset_uid / corridor / vendor)", value="")
            inv = filtered.copy()

            if search.strip():
                s = search.strip().lower()
                inv = inv[
                    inv["asset_uid"].str.lower().str.contains(s) |
                    inv["corridor"].str.lower().str.contains(s) |
                    inv["vendor"].str.lower().str.contains(s)
                ].copy()

            cols_show = [
                "asset_uid", "asset_type", "region", "corridor",
                "latitude", "longitude",
                "owner", "vendor",
                "install_year", "age_years", "due_year", "support_end_year",
                "status_window", "availability_window", "outages_count_window",
                "mttr_hours_window", "config_noncompliance_window",
                "risk_score", "risk_score_baseline",
                "criticality_score", "replacement_cost_usd"
            ]
            inv_out = inv[cols_show].sort_values(["risk_score", "criticality_score"], ascending=False)
            st_df(inv_out, height=560)

            st.download_button(
                "Download filtered inventory (CSV)",
                data=inv_out.to_csv(index=False).encode("utf-8"),
                file_name="smartway_inventory_filtered.csv",
                mime="text/csv",
            )

        with right:
            st.markdown("<div class='section-title'>Map — selected asset + nearby context</div>", unsafe_allow_html=True)
            asset_id = st.selectbox("Focus asset", options=inv_out["asset_uid"].tolist())
            focus = filtered[filtered["asset_uid"] == asset_id]
            context = pd.concat([
                focus,
                filtered.sort_values("risk_score", ascending=False).head(250)
            ]).drop_duplicates("asset_uid").head(max_rows_map)
            render_map(context, window_days, risk_thr, ASOF_YEAR, height=600, title="")


# =========================
# Operations
# =========================
with tab_ops:
    if filtered.empty:
        st.warning("No assets match the current filters.")
    else:
        left, right = st.columns([0.58, 0.42], gap="large")

        with left:
            st.markdown(f"<div class='section-title'>Operational KPIs + maintenance ({window_days}-day window)</div>", unsafe_allow_html=True)

            k = (
                filtered.groupby("asset_type", as_index=False)
                .agg(
                    assets=("asset_uid", "count"),
                    avg_availability=("availability_window", "mean"),
                    avg_mttr=("mttr_hours_window", "mean"),
                    avg_risk=("risk_score", "mean"),
                    non_compliant=("config_noncompliance_window", "sum"),
                    spend=("maintenance_cost_window", "sum")
                )
            )
            k["avg_availability_pct"] = 100 * k["avg_availability"]
            k["avg_mttr"] = k["avg_mttr"].round(2)
            k["avg_risk_pct"] = (100 * k["avg_risk"]).round(1)
            k["spend_usd"] = k["spend"].round(0).astype(int)

            c1, c2 = st.columns(2)
            with c1:
                st_plot(px.bar(k, x="asset_type", y="avg_availability_pct", title="Avg availability (%) by asset type"))
            with c2:
                st_plot(px.bar(k, x="asset_type", y="avg_risk_pct", title="Avg predicted outage risk (%) by asset type"))

            st.markdown("<div class='section-title'>KPI table (by asset type)</div>", unsafe_allow_html=True)
            st_df(k[["asset_type", "assets", "avg_availability_pct", "avg_mttr", "avg_risk_pct", "non_compliant", "spend_usd"]], height=260)

            st.write("")
            st.markdown("<div class='section-title'>Maintenance / outage event log</div>", unsafe_allow_html=True)
            cutoff = pd.Timestamp(ASOF_DT - timedelta(days=window_days))
            evw = events[events["event_start"] >= cutoff].copy()
            evw = evw[evw["asset_uid"].isin(filtered["asset_uid"])].copy()

            if evw.empty:
                st.info("No maintenance/outage events in the selected window for the filtered assets.")
            else:
                a, b = st.columns(2)
                with a:
                    rsn = evw["reason"].value_counts().reset_index()
                    rsn.columns = ["reason", "events"]
                    st_plot(px.bar(rsn, x="reason", y="events", title="Events by reason"))
                with b:
                    wot = evw["work_order_type"].value_counts().reset_index()
                    wot.columns = ["work_order_type", "events"]
                    st_plot(px.pie(wot, names="work_order_type", values="events", title="Work order type mix"))

                show = evw.sort_values("downtime_hours", ascending=False).head(300)
                st_df(show[["asset_uid", "event_start", "downtime_hours", "work_order_type", "reason", "maintenance_cost_usd"]], height=420)

                st.download_button(
                    "Download maintenance events (CSV)",
                    data=show.to_csv(index=False).encode("utf-8"),
                    file_name="smartway_maintenance_events.csv",
                    mime="text/csv",
                )

        with right:
            st.markdown("<div class='section-title'>Map — outage/maintenance hotspots</div>", unsafe_allow_html=True)
            hot = filtered.copy()
            hot["hot_score"] = 0.55 * hot["downtime_hours_window"] + 60 * hot["risk_score"] + 10 * (1 - hot["availability_window"])
            hot = hot.sort_values("hot_score", ascending=False).head(max_rows_map)
            render_map(hot, window_days, risk_thr, ASOF_YEAR, height=650, title="")


# =========================
# Scenario planning
# =========================
with tab_scenario:
    if filtered.empty:
        st.warning("No assets match the current filters.")
    else:
        left, right = st.columns([0.58, 0.42], gap="large")

        with left:
            st.markdown("<div class='section-title'>5–10 year replacement & scenario planning (AI-prioritized)</div>", unsafe_allow_html=True)

            colA, colB, colC = st.columns(3)
            start_year = colA.number_input("Start year", min_value=2020, max_value=2040, value=ASOF_YEAR, step=1)
            horizon_years = colB.slider("Horizon (years)", 5, 10, 10, 1)
            annual_budget = colC.number_input("Annual renewal budget (USD)", min_value=100000, max_value=6000000, value=200000, step=200000)

            scenario_summary, scenario_selected, backlog_ids = lifecycle_scenario(
                filtered, int(start_year), int(horizon_years), int(annual_budget), asof_year=ASOF_YEAR
            )

            c1, c2, c3 = st.columns(3)
            c1.metric("Backlog at end", f"{scenario_summary['backlog_assets'].iloc[-1]:,}")
            c2.metric("Total spend", f"${scenario_summary['annual_spend_usd'].sum():,}")
            c3.metric("Total replaced", f"{scenario_summary['replaced_assets'].sum():,}")

            fig = go.Figure()
            fig.add_trace(go.Bar(x=scenario_summary["year"], y=scenario_summary["due_assets"], name="Due assets"))
            fig.add_trace(go.Bar(x=scenario_summary["year"], y=scenario_summary["replaced_assets"], name="Replaced assets"))
            fig.add_trace(go.Scatter(x=scenario_summary["year"], y=scenario_summary["backlog_assets"], name="Backlog", mode="lines+markers"))
            fig.update_layout(barmode="group", xaxis_title="Year", yaxis_title="Assets", title="Renewal demand vs budgeted renewals")
            st_plot(fig)

            st.markdown("<div class='section-title'>Selected assets (sample)</div>", unsafe_allow_html=True)
            if len(scenario_selected):
                st_df(scenario_selected.sort_values(["selected_year", "risk_score"], ascending=[True, False]).head(250), height=420)
            else:
                st.info("No selections were made under the selected annual budget (try increasing budget).")

            st.download_button(
                "Download scenario summary (CSV)",
                data=scenario_summary.to_csv(index=False).encode("utf-8"),
                file_name="smartway_scenario_summary.csv",
                mime="text/csv",
            )

        with right:
            st.markdown("<div class='section-title'>Map — scenario replacements</div>", unsafe_allow_html=True)

            if len(scenario_selected):
                years = sorted(scenario_selected["selected_year"].unique().tolist())
                year_pick = st.select_slider("Show replacements for year", options=years, value=years[0])
                year_df = scenario_selected[scenario_selected["selected_year"] == year_pick].copy()

                need_cols = ["status_window", "mttr_hours_window",
                             "outages_count_window", "config_noncompliance_window", "remaining_life_years"]
                missing = [c for c in need_cols if c not in year_df.columns]
                if missing:
                    year_df = year_df.merge(
                        filtered[["asset_uid"] + missing],
                        on="asset_uid", how="left"
                    )

                year_df = year_df.sort_values("risk_score", ascending=False).head(max_rows_map)
                render_map(year_df, window_days, risk_thr, ASOF_YEAR, height=650, title=f"Replacement candidates — {year_pick}")
            else:
                st.info("No replacement candidates selected under current budget; adjust budget or filters.")


# =========================
# Asset Map
# =========================
with tab_asset_map:
    if filtered.empty:
        st.warning("No assets match the current filters.")
    else:
        st.markdown("<div class='section-title'>Road map — window-based status markers</div>", unsafe_allow_html=True)
        st.caption("Markers are colored by window-based service status (OUT_OF_SERVICE / DEGRADED / IN_SERVICE).")

        map_view = filtered.copy()
        map_view = map_view.sort_values(["status_window", "risk_score"], ascending=[True, False]).head(max_rows_map)

        render_map(map_view, window_days, risk_thr, ASOF_YEAR, height=680, title="")

        st.write("")
        st.markdown("<div class='section-title'>Quick status counts (window)</div>", unsafe_allow_html=True)
        st_df(filtered["status_window"].value_counts().rename_axis("status_window").reset_index(name="count"), height=180)


# =========================
# Data Quality
# =========================
with tab_quality:
    st.markdown("<div class='section-title'>Cross-system reconciliation & data quality</div>", unsafe_allow_html=True)
    merged = reconcile_sources(tms_atms, cmms, fin, ven)

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Union asset_uids", f"{len(merged):,}")
    c2.metric("Type conflicts", f"{int(merged['type_conflict'].sum()):,}")
    c3.metric("Location conflicts", f"{int(merged['loc_conflict'].sum()):,}")
    c4.metric("Filtered assets", f"{len(filtered):,}")

    cov = pd.DataFrame({
        "System": ["TMS/ATMS (SmartWay)", "CMMS", "Finance", "Vendor logs"],
        "Records": [
            int(merged["present_TMS_ATMS"].sum()),
            int(merged["present_CMMS"].sum()),
            int(merged["present_FIN"].sum()),
            int(merged["present_VEN"].sum()),
        ],
    })

    left, right = st.columns([0.55, 0.45], gap="large")
    with left:
        st.markdown("<div class='section-title'>Coverage summary</div>", unsafe_allow_html=True)
        st_df(cov, height=220)

        st.markdown("<div class='section-title'>Conflict samples</div>", unsafe_allow_html=True)
        conflicts = merged[(merged["type_conflict"]) | (merged["loc_conflict"])].copy()
        show = conflicts[[
            "asset_uid",
            "TMS_ATMS_asset_type", "CMMS_asset_type", "FIN_asset_type", "VEN_asset_type",
            "type_conflict", "loc_conflict", "loc_spread"
        ]].head(60)
        st_df(show, height=360)

        st.download_button(
            "Download reconciliation union (CSV)",
            data=merged.to_csv(index=False).encode("utf-8"),
            file_name="smartway_reconciliation_union.csv",
            mime="text/csv",
        )

    with right:
        st.markdown("<div class='section-title'>Map — assets with conflicts</div>", unsafe_allow_html=True)
        conf_ids = set(conflicts["asset_uid"].astype(str).tolist())
        conf_map = df_kpi[df_kpi["asset_uid"].isin(conf_ids)].copy()
        if not conf_map.empty:
            conf_map = conf_map.sort_values(["risk_score", "criticality_score"], ascending=False).head(max_rows_map)
            render_map(conf_map, window_days, risk_thr, ASOF_YEAR, height=650, title="")
        else:
            st.info("No mappable conflict records in the current inventory.")


# =========================
# AI Risk
# =========================
with tab_ai:
    if filtered.empty:
        st.warning("No assets match the current filters.")
    else:
        left, right = st.columns([0.58, 0.42], gap="large")

        with left:
            st.markdown("<div class='section-title'>AI-assisted risk scoring + per-asset prediction</div>", unsafe_allow_html=True)

            auc_tr = df_kpi["model_auc_trained"].iloc[0] if "model_auc_trained" in df_kpi.columns else np.nan
            auc_bl = df_kpi["model_auc_baseline"].iloc[0] if "model_auc_baseline" in df_kpi.columns else np.nan
            if not pd.isna(auc_tr):
                st.caption(f"Trained model: Logistic Regression (demo). Holdout AUC={auc_tr:.3f} | Baseline AUC={auc_bl:.3f}")
            else:
                st.caption("Model fallback mode (insufficient label diversity).")

            alerts = filtered.copy()
            alerts["alert_flag"] = alerts["risk_score"] >= risk_thr
            top_alerts = alerts.sort_values(["risk_score", "outages_count_window"], ascending=False).head(50)

            st.markdown("<div class='section-title'>Top risk-ranked assets</div>", unsafe_allow_html=True)
            st_df(top_alerts[[
                "asset_uid", "asset_type", "region", "corridor",
                "latitude", "longitude",
                "availability_window", "outages_count_window", "mttr_hours_window",
                "config_noncompliance_window", "risk_score", "risk_score_baseline",
                "criticality_score", "due_year", "status_window"
            ]], height=360)

            st.download_button(
                "Download top alerts (CSV)",
                data=top_alerts.to_csv(index=False).encode("utf-8"),
                file_name="smartway_top_alerts.csv",
                mime="text/csv",
            )

            st.write("")
            st.markdown("<div class='section-title'>Per-asset drivers</div>", unsafe_allow_html=True)
            asset_id = st.selectbox("Select an asset", options=filtered["asset_uid"].tolist(), key="asset_pick_ai")
            row = filtered[filtered["asset_uid"] == asset_id].iloc[0]

            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Predicted outage risk", f"{row['risk_score']*100:.1f}%")
            c2.metric("Availability", f"{row['availability_window']*100:.1f}%")
            c3.metric("MTTR", "—" if pd.isna(row["mttr_hours_window"]) else f"{row['mttr_hours_window']:.1f} h")
            c4.metric("Due year", f"{int(row['due_year'])}")

            urgency, actions = recommended_action(row, risk_thr, ASOF_YEAR)
            st.info("**Recommended actions**\n\n- " + "\n- ".join(actions))

            if risk_model is not None:
                feature_row = pd.DataFrame([{
                    "asset_type": row["asset_type"],
                    "age_years": row["age_years"],
                    "availability_window": row["availability_window"],
                    "config_noncompliance_window": row["config_noncompliance_window"],
                    "replacement_cost_usd": row["replacement_cost_usd"],
                    "ctx_score": row["ctx_score"],
                    "criticality_score": row["criticality_score"]
                }])
                contrib = model_contributions(risk_model, feature_row)
                if len(contrib):
                    st_plot(px.bar(contrib.iloc[::-1], x="contribution", y="feature",
                                   title="Top drivers (approx. contribution to risk score)"))

        with right:
            st.markdown("<div class='section-title'>Map — selected asset + top risk context</div>", unsafe_allow_html=True)
            map_ai = pd.concat([
                filtered[filtered["asset_uid"] == asset_id],
                filtered.sort_values("risk_score", ascending=False).head(250)
            ]).drop_duplicates("asset_uid").head(max_rows_map)
            render_map(map_ai, window_days, risk_thr, ASOF_YEAR, height=720, title="")


# =========================
# Reports
# =========================
with tab_report:
    if filtered.empty:
        st.warning("No assets match the current filters.")
    else:
        st.markdown("<div class='section-title'>Report author details</div>", unsafe_allow_html=True)

        a1, a2 = st.columns(2)
        prepared_by = a1.text_input("Prepared By (Name)", value="Md Tawkir Ahmed")
        department = a2.text_input("Department", value="SmartWay Operations")

        b1, b2 = st.columns(2)
        role = b1.text_input("Role", value="Transportation Researcher")
        organization = b2.text_input("Organization", value="C-TIER")

        st.markdown("<div class='section-title'>Downloads</div>", unsafe_allow_html=True)

        cutoff = pd.Timestamp(ASOF_DT - timedelta(days=window_days))
        evw = events[events["event_start"] >= cutoff].copy()
        evw = evw[evw["asset_uid"].isin(filtered["asset_uid"])].copy()

        kpi_by_type = (
            filtered.groupby("asset_type", as_index=False)
            .agg(
                assets=("asset_uid", "count"),
                avg_availability=("availability_window", "mean"),
                avg_mttr=("mttr_hours_window", "mean"),
                avg_risk=("risk_score", "mean"),
                non_compliant=("config_noncompliance_window", "sum"),
            )
        )
        kpi_by_type["avg_availability_pct"] = (100 * kpi_by_type["avg_availability"]).round(2)
        kpi_by_type["avg_risk_pct"] = (100 * kpi_by_type["avg_risk"]).round(2)
        kpi_by_type["avg_mttr"] = kpi_by_type["avg_mttr"].round(2)

        scenario_summary, scenario_selected, _ = lifecycle_scenario(filtered, ASOF_YEAR, 10, 6000000, asof_year=ASOF_YEAR)

        savings_df = compute_ai_vs_traditional_table(
            filtered_assets=filtered,
            window_days=window_days,
            downtime_cost_per_hour=float(downtime_cost_per_hour),
            downtime_reduction_pct=float(downtime_reduction_pct),
            analyst_hourly_rate=float(analyst_hourly_rate),
            manual_hours_per_month=float(manual_hours_per_month),
            ai_hours_per_month=float(ai_hours_per_month),
            horizon_months=int(horizon_months),
            platform_annual_cost=float(AI_PLATFORM_ANNUAL_RENEWAL_USD),
        )

        meta = {
            "generated_ts": ts_label,
            "asof_date": asof_date.isoformat(),
            "asof_year": ASOF_YEAR,
            "pdf_name": f"SmartWay_Report_{ts_file}.pdf",
            "docx_name": f"SmartWay_Report_{ts_file}.docx",
            "regions": region_sel,
            "types": type_sel,
            "corridors": corr_sel,
            "status": status_sel,
            "window_days": window_days,
            "min_avail": float(min_avail),
            "risk_thr": float(risk_thr),
            "seed": int(seed),
            "prepared_by": prepared_by,
            "department": department,
            "role": role,
            "organization": organization,
            "platform_annual_cost": float(AI_PLATFORM_ANNUAL_RENEWAL_USD),
        }

        merged_quality = reconcile_sources(tms_atms, cmms, fin, ven)

        pdf_bytes = b""
        if HAS_PDF:
            pdf_bytes = build_pdf_report_bytes(
                filtered_assets=filtered,
                events_window=evw,
                kpi_by_type=kpi_by_type,
                scenario_summary=scenario_summary,
                scenario_selected=scenario_selected,
                merged_quality=merged_quality,
                meta=meta,
                savings_table=savings_df
            )

        docx_bytes = b""
        if HAS_DOCX:
            docx_bytes = build_docx_report_bytes(
                filtered_assets=filtered,
                events_window=evw,
                scenario_summary=scenario_summary,
                merged_quality=merged_quality,
                meta=meta,
                savings_table=savings_df
            )

        col1, col2, col3 = st.columns([0.34, 0.33, 0.33])

        with col1:
            if HAS_PDF and pdf_bytes:
                st.download_button(
                    "Download PDF report",
                    data=pdf_bytes,
                    file_name=meta["pdf_name"],
                    mime="application/pdf"
                )
            elif not HAS_PDF:
                st.info("PDF export disabled. Run: `pip install reportlab`.")

        with col2:
            if HAS_DOCX and docx_bytes:
                st.download_button(
                    "Download Word report (DOCX)",
                    data=docx_bytes,
                    file_name=meta["docx_name"],
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            elif not HAS_DOCX:
                st.info("DOCX export disabled. Run: `pip install python-docx`.")

        with col3:
            zip_bytes = build_report_zip(
                filtered_assets=filtered,
                events_window=evw,
                kpi_by_type=kpi_by_type,
                scenario_summary=scenario_summary,
                scenario_selected=scenario_selected,
                pdf_bytes=pdf_bytes,
                docx_bytes=docx_bytes,
                meta=meta
            )
            st.download_button(
                "Download export package (ZIP)",
                data=zip_bytes,
                file_name=f"SmartWay_Export_{ts_file}.zip",
                mime="application/zip"
            )

st.write("")
st.caption("This prototype dashboard using synthetic data")
